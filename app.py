from flask import Flask, render_template, request, send_file, jsonify
import docx2pdf
from docx2pdf import convert
from docx import Document
from pdf2docx import parse
import os
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename
import asyncio
import aiofiles
#Image 
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Inches
import docx2txt
import tempfile
#Merging
import io
import PyPDF2
from PyPDF2 import PdfFileWriter, PdfFileReader
import fitz #pip install PyMuPDF

import tkinter as to 
import tkinter.ttk as ttk
from  tkinter.filedialog import askopenfile
from tkinter.messagebox  import showinfo

app = Flask(__name__)

ALLOWED_EXTENSIONS = {'docx', 'pdf'}

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
           
async def parse_async(file_path):
    with open(file_path, 'rb') as f:
        async with aiofiles.tempfile.TemporaryFile() as tmp:
            async for chunk in aiofiles.threadpool.binary_reader(f):
                await tmp.write(chunk)
            return await parse(tmp.name)

def convert(input_path, output_path):
    print("Converting {} to {}".format(input_path, output_path))
    try:
        if output_path.endswith('.pdf'):
            # Convert docx to pdf
            doc = Document(input_path)
            doc.save(output_path)
        elif output_path.endswith('.docx'):
            # Convert pdf to docx
            pdf = PdfFileReader(open(input_path, 'rb'))
            doc = Document()
            for page in range(pdf.getNumPages()):
                doc.add_paragraph(pdf.getPage(page).extractText())
            doc.save(output_path)
        print("Conversion completed successfully")
    except Exception as e:
        print("Error during conversion: {}".format(e))
        
def convert_to_image(file):
    # Create a BytesIO buffer to hold the image data
    image_buffer = BytesIO()
    
    # Check the file type and convert as needed
    if file.content_type.startswith('image/'):
        # File is already an image, so just return it as a BytesIO object
        return BytesIO(file.read())
    elif file.content_type == 'application/pdf':
        # File is a PDF, so convert the first page to an image
        pdf_buffer = BytesIO(file.read())
        pdf_reader = PyPDF2.PdfReader(pdf_buffer)
        page = pdf_reader.pages[0]
        page_data = page['/Resources']['/XObject'].get_object()['/Im0'].getData()
        image = Image.open(BytesIO(page_data))
        return image
                  
def convert_to_pdf(file):
    # Create a BytesIO buffer to hold the PDF data
    pdf_buffer = BytesIO()
    
    # Check the file type and convert as needed
    if file.content_type == 'application/pdf':
        # File is already a PDF, so just return it as a BytesIO object
        return BytesIO(file.read())
    elif file.content_type.startswith('image/'):
        # File is an image, so convert it to a PDF
        image = Image.open(file)
        image_pdf = image.convert('RGB')
        image_pdf.save(pdf_buffer, 'PDF')
    elif file.content_type == 'application/msword':
        # File is a Word document, so convert it to a PDF
        os.makedirs('/tmp/word_files', exist_ok=True)
        word_file_path = os.path.join('/tmp/word_files', file.filename)
        file.save(word_file_path)
        os.system(f'libreoffice --convert-to pdf {word_file_path} --outdir /tmp/word_files')
        pdf_file_path = os.path.join('/tmp/word_files', f'{os.path.splitext(file.filename)[0]}.pdf')
        with open(pdf_file_path, 'rb') as f:
            pdf_buffer.write(f.read())
        os.remove(pdf_file_path)
        os.remove(word_file_path)
    else:
        # Unsupported file type
        raise Exception(f'Unsupported file type: {file.content_type}')
    
    pdf_buffer.seek(0)  # Reset the buffer position to the beginning
    return pdf_buffer

# route to convert docx to PDF3
@app.route('/docx_to_pdf', methods=['POST'])
def docx_to_pdf():
    if 'file' not in request.files:
        return 'No file uploaded', 400

    file = request.files['file']

    if not file.filename.lower().endswith('.docx'):
        return 'Only .docx files are supported', 400

    docx_filename = file.filename
    pdf_filename = os.path.splitext(docx_filename)[0] + '.pdf'

    file.save(docx_filename)

    try:
        convert(docx_filename, pdf_filename)
    except Exception as e:
        return str(e), 500

    if os.path.exists(pdf_filename):
        return send_file(pdf_filename)
    else:
        return 'File does not exist', 500
    
# route to convert PDF to Docx
@app.route('/pdf_to_docx', methods=['POST'])
def pdf_to_docx():
    file = request.files['file']
    file.save(file.filename)
    parse(file.filename, file.filename[:-4] + '.docx')
    return send_file(file.filename[:-4] + '.docx')

# route for the merging PDFs
@app.route('/merge_pdfs', methods=['POST'])
def merge_pdfs():
    # Get list of uploaded files
    files = request.files.getlist('files')
    
    # Create PdfFileMerger object
    merger = PyPDF2.PdfMerger()

    # Loop through each file
    for file in files:
        # Read PDF file
        pdf = PyPDF2.PdfReader(file)
        # Add PDF to merger object
        merger.append(pdf)

    # Write merged PDF to memory
    merger.write('merged.pdf')
    
    # Return merged PDF as attachment
    return send_file('merged.pdf')


# route for the image to PDF conversion page
@app.route('/image-to-pdf', methods=['GET', 'POST'])
def image_to_pdf():
    file = request.files['file']
    pdf_buffer = convert_to_pdf(file)
    return send_file(pdf_buffer, mimetype='application/pdf', as_attachment=True,
                     download_name=f'converted_files.pdf')

#route for the PDF to image conversion page
@app.route('/pdf-to-image', methods=['POST'])
def pdf_to_image():
    # Check if a file was uploaded
    if 'file' not in request.files:
        return 'No file uploaded', 400
    
    # Get the uploaded file
    file = request.files['file']
    
    # Check if the file is a PDF
    if file.filename.endswith('.pdf'):
        pdf_file = fitz.open(stream=io.BytesIO(file.read()))
        
        # Select the first page of the PDF
        page = pdf_file[0]
        
        # Create a pixmap object from the page
        pixmap = page.get_pixmap()
        image_bytes = pixmap.tobytes()
        
        # Create a file-like object from the image bytes
        image_file = io.BytesIO(image_bytes)
        
        # Return the image file
        return send_file(image_file, mimetype='image/png')
    else:
        return 'File is not a PDF', 400
    
#route for the image to Docx conversion page
@app.route('/image-to-docx', methods=['POST'])
def convert_Image_to_Docx():
     # Get the image file from the request
    file = request.files['file']
    # Open the image file
    img = Image.open(file)
    #Convert the image to PDF and save it to disk
    if img.mode == 'RGBA':
     img = img.convert('RGB')
    pdf_file = open('output.pdf', 'wb')
    img.save(pdf_file, 'PDF')
    pdf_file.close()
    parse('output.pdf', 'output.docx')
    return send_file('output.docx')

@app.route('/docx_to_image', methods=['POST'])
def docxToImage():
    dir_path = '/Users/moon/Downloads/upload/' #Put your directory path here
    # Get the input file and output format from the request
    try:
     image_files = [filename for filename in os.listdir(dir_path) if filename.endswith('.png') or filename.endswith('.jpg')or filename.endswith('.jpeg')]
     for image_file in image_files:
       os.remove(dir_path+image_file)
     input_file = request.files['file']
     input_file.save(input_file.filename)
     text = docx2txt.process(input_file.filename, dir_path)
     image_files = [filename for filename in os.listdir(dir_path) if filename.endswith('.png') or filename.endswith('.jpg')or filename.endswith('.jpeg')]
     largest_file = ''
     largest_file_size = 0
     for file_name in image_files:
      file_path = os.path.join(dir_path, file_name)
      file_size = os.path.getsize(file_path)
    # Check if file size is larger than current largest file
      if file_size > largest_file_size:
        largest_file = file_name
        largest_file_size = file_size
     return send_file(dir_path+largest_file, mimetype='image/png')
    except Exception as e:
        error_msg = str(e)
        return jsonify({'error': error_msg})
    
        


